from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models import ExportFile, ExportStatus, ExportType, Explanation, Report, ReportSection, Requirement
from app.services.reports import build_report_explanations_payload
from app.services.storage import storage


def export_report_docx(db: Session, report: Report, created_by_id: str | None) -> ExportFile:
    file_name = f"{report.title}_report.docx".replace(" ", "_")
    path = Path(storage.create_export_path(report.organization_id, file_name))
    document = DocxDocument()
    document.add_heading(report.title, level=1)
    for section in db.scalars(select(ReportSection).where(ReportSection.report_id == report.id).order_by(ReportSection.order_number)):
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.content)
    document.save(path)
    export = ExportFile(
        organization_id=report.organization_id,
        report_id=report.id,
        created_by_id=created_by_id,
        export_type=ExportType.docx,
        file_name=file_name,
        storage_path=str(path),
        status=ExportStatus.ready,
    )
    db.add(export)
    db.commit()
    db.refresh(export)
    return export


def export_matrix_xlsx(db: Session, report: Report, created_by_id: str | None) -> ExportFile:
    file_name = f"{report.title}_matrix.xlsx".replace(" ", "_")
    path = Path(storage.create_export_path(report.organization_id, file_name))
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Matrix"
    sheet.append(["Requirement ID", "Category", "Title", "Status", "Confidence", "Risk", "Applicability"])
    for requirement in db.scalars(select(Requirement).where(Requirement.report_id == report.id).order_by(Requirement.created_at)):
        sheet.append(
            [
                requirement.id,
                requirement.category,
                requirement.title,
                requirement.status.value,
                requirement.confidence_score,
                requirement.risk_level.value,
                requirement.applicability_status.value,
            ]
        )
    workbook.save(path)
    export = ExportFile(
        organization_id=report.organization_id,
        report_id=report.id,
        created_by_id=created_by_id,
        export_type=ExportType.matrix,
        file_name=file_name,
        storage_path=str(path),
        status=ExportStatus.ready,
    )
    db.add(export)
    db.commit()
    db.refresh(export)
    return export


def export_explanations_html(db: Session, report: Report, created_by_id: str | None) -> ExportFile:
    file_name = f"{report.title}_explanations.html".replace(" ", "_")
    path = Path(storage.create_export_path(report.organization_id, file_name))
    explanations_payload = build_report_explanations_payload(db, report)
    blocks: list[str] = []
    for item in explanations_payload:
        logic_lines = "".join(f"<li>{line}</li>" for line in item["logic"])
        recommended = f"<p><strong>Рекомендация:</strong> {item['recommended_action']}</p>" if item["recommended_action"] else ""
        blocks.append(
            f"""
            <section class="explanation-card">
              <h2>{item['conclusion']}</h2>
              <p>{item['explanation_text']}</p>
              <p><strong>Уверенность:</strong> {item['confidence_score']}</p>
              <p><strong>Риск:</strong> {item['risk_level']}</p>
              <ul>{logic_lines}</ul>
              {recommended}
            </section>
            """
        )
    html = f"""
    <!doctype html>
    <html lang="ru">
      <head>
        <meta charset="utf-8" />
        <title>XAI explanations - {report.title}</title>
        <style>
          body {{ font-family: Arial, sans-serif; margin: 2rem; color: #1d2430; }}
          .explanation-card {{ border: 1px solid #d7dde4; border-radius: 12px; padding: 1rem 1.25rem; margin-bottom: 1rem; }}
          h1 {{ margin-bottom: 1.5rem; }}
          h2 {{ font-size: 1.1rem; margin: 0 0 0.75rem; }}
        </style>
      </head>
      <body>
        <h1>XAI-объяснения: {report.title}</h1>
        {''.join(blocks) if blocks else '<p>Объяснения пока не сформированы.</p>'}
      </body>
    </html>
    """
    path.write_text(html, encoding="utf-8")
    export = ExportFile(
        organization_id=report.organization_id,
        report_id=report.id,
        created_by_id=created_by_id,
        export_type=ExportType.explanations,
        file_name=file_name,
        storage_path=str(path),
        status=ExportStatus.ready,
    )
    db.add(export)
    db.commit()
    db.refresh(export)
    return export


def export_evidence_package(db: Session, report: Report, created_by_id: str | None) -> ExportFile:
    docx_export = export_report_docx(db, report, created_by_id)
    matrix_export = export_matrix_xlsx(db, report, created_by_id)
    explanations_export = export_explanations_html(db, report, created_by_id)
    file_name = f"{report.title}_package.zip".replace(" ", "_")
    path = Path(storage.create_export_path(report.organization_id, file_name))

    explanations_payload = build_report_explanations_payload(db, report)

    explanations_path = path.with_suffix(".json")
    explanations_path.write_text(json.dumps(explanations_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    with zipfile.ZipFile(path, "w") as archive:
        archive.write(docx_export.storage_path, arcname=Path(docx_export.storage_path).name)
        archive.write(matrix_export.storage_path, arcname=Path(matrix_export.storage_path).name)
        archive.write(explanations_export.storage_path, arcname=Path(explanations_export.storage_path).name)
        archive.write(explanations_path, arcname="explanations.json")

    export = ExportFile(
        organization_id=report.organization_id,
        report_id=report.id,
        created_by_id=created_by_id,
        export_type=ExportType.package,
        file_name=file_name,
        storage_path=str(path),
        status=ExportStatus.ready,
    )
    db.add(export)
    db.commit()
    db.refresh(export)
    return export
