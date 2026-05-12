from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from io import StringIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from app.integrations.ocr import OCRError, get_ocr_provider
from app.models import FragmentType


@dataclass
class FragmentSeed:
    text: str
    fragment_type: FragmentType
    page_number: int | None = None
    sheet_name: str | None = None
    row_start: int | None = None
    row_end: int | None = None
    paragraph_number: int | None = None


@dataclass
class ExtractionResult:
    text: str
    fragments: list[FragmentSeed]
    requires_review: bool = False
    page_count: int | None = None


def _page_fragment_from_text(text: str, *, page_number: int) -> FragmentSeed:
    return FragmentSeed(text=text.strip(), fragment_type=FragmentType.page, page_number=page_number)


def _extract_text_file(path: Path) -> ExtractionResult:
    text = path.read_text(encoding="utf-8", errors="ignore")
    fragments = [
        FragmentSeed(text=chunk.strip(), fragment_type=FragmentType.paragraph, paragraph_number=index + 1)
        for index, chunk in enumerate(text.splitlines())
        if chunk.strip()
    ]
    return ExtractionResult(text=text, fragments=fragments, page_count=len(fragments) or 1)


def _extract_json(path: Path) -> ExtractionResult:
    raw = json.loads(path.read_text(encoding="utf-8"))
    pretty = json.dumps(raw, ensure_ascii=False, indent=2)
    fragments = [
        FragmentSeed(text=chunk.strip(), fragment_type=FragmentType.paragraph, paragraph_number=index + 1)
        for index, chunk in enumerate(pretty.splitlines())
        if chunk.strip()
    ]
    return ExtractionResult(text=pretty, fragments=fragments, page_count=1)


def _extract_csv(path: Path) -> ExtractionResult:
    reader = csv.reader(StringIO(path.read_text(encoding="utf-8", errors="ignore")))
    fragments: list[FragmentSeed] = []
    lines: list[str] = []
    for index, row in enumerate(reader):
        if not any(cell.strip() for cell in row):
            continue
        line = " | ".join(cell.strip() for cell in row)
        lines.append(line)
        fragments.append(FragmentSeed(text=line, fragment_type=FragmentType.sheet_row, row_start=index + 1, row_end=index + 1))
    return ExtractionResult(text="\n".join(lines), fragments=fragments, page_count=1)


def _extract_xlsx(path: Path) -> ExtractionResult:
    workbook = load_workbook(path, data_only=True)
    lines: list[str] = []
    fragments: list[FragmentSeed] = []
    for sheet in workbook.worksheets:
        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if not values:
                continue
            line = " | ".join(values)
            lines.append(f"[{sheet.title}] {line}")
            fragments.append(
                FragmentSeed(
                    text=line,
                    fragment_type=FragmentType.sheet_row,
                    sheet_name=sheet.title,
                    row_start=row_index,
                    row_end=row_index,
                )
            )
    return ExtractionResult(text="\n".join(lines), fragments=fragments, page_count=len(workbook.worksheets))


def _extract_docx(path: Path) -> ExtractionResult:
    document = DocxDocument(path)
    fragments: list[FragmentSeed] = []
    lines: list[str] = []
    paragraph_index = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        paragraph_index += 1
        lines.append(text)
        fragments.append(FragmentSeed(text=text, fragment_type=FragmentType.paragraph, paragraph_number=paragraph_index))

    for table in document.tables:
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if not cells:
                continue
            line = " | ".join(cells)
            lines.append(line)
            fragments.append(FragmentSeed(text=line, fragment_type=FragmentType.table_row, row_start=row_index, row_end=row_index))

    return ExtractionResult(text="\n".join(lines), fragments=fragments, page_count=max(paragraph_index, 1))


def _extract_pdf(path: Path) -> ExtractionResult:
    reader = PdfReader(str(path))
    fragments: list[FragmentSeed] = []
    lines: list[str] = []
    requires_review = False
    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            text = _extract_pdf_page_with_ocr(path, page_index)
        if not text:
            requires_review = True
            continue
        lines.append(text)
        fragments.append(_page_fragment_from_text(text, page_number=page_index))
    return ExtractionResult(text="\n\n".join(lines), fragments=fragments, requires_review=requires_review, page_count=len(reader.pages))


def _extract_image(path: Path) -> ExtractionResult:
    try:
        ocr_result = get_ocr_provider().extract_text(path)
    except OCRError:
        return ExtractionResult(text="", fragments=[], requires_review=True, page_count=1)

    fragments = [_page_fragment_from_text(ocr_result.text, page_number=1)] if ocr_result.text.strip() else []
    return ExtractionResult(text=ocr_result.text, fragments=fragments, requires_review=not bool(fragments), page_count=1)


def _render_pdf_page_for_ocr(path: Path, page_number: int) -> Any:
    try:
        import pypdfium2 as pdfium  # type: ignore
    except Exception as exc:  # pragma: no cover - optional dependency path
        raise OCRError("PDF OCR rendering dependencies are not installed") from exc

    pdf = None
    page = None
    try:
        pdf = pdfium.PdfDocument(str(path))
        page = pdf[page_number - 1]
        pil_image = page.render(scale=2.0).to_pil()
        return pil_image
    except Exception as exc:  # pragma: no cover - defensive
        raise OCRError(f"Unable to render PDF page {page_number} for OCR") from exc
    finally:  # pragma: no branch - defensive cleanup
        if page is not None and hasattr(page, "close"):
            page.close()
        if pdf is not None and hasattr(pdf, "close"):
            pdf.close()


def _extract_pdf_page_with_ocr(path: Path, page_number: int) -> str:
    try:
        image = _render_pdf_page_for_ocr(path, page_number)
        ocr_result = get_ocr_provider().extract_image_object(image, source_name=f"{path.name}#page={page_number}")
    except OCRError:
        return ""
    return ocr_result.text.strip()


def extract_document(path_str: str) -> ExtractionResult:
    path = Path(path_str)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".log"}:
        return _extract_text_file(path)
    if suffix == ".json":
        return _extract_json(path)
    if suffix == ".csv":
        return _extract_csv(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}:
        return _extract_image(path)
    raise ValueError(f"Unsupported document format: {suffix}")
