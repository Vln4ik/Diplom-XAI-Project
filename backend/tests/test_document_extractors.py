from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from pypdf import PdfWriter

from app.processors.documents import extract_document


def test_extract_text_json_csv_docx_xlsx_pdf(tmp_path: Path):
    txt_path = tmp_path / "sample.txt"
    txt_path.write_text("Лицензия\nКадровый состав\n", encoding="utf-8")

    json_path = tmp_path / "profile.json"
    json_path.write_text(
        json.dumps({"website": "https://college.example.edu", "accreditation": True}, ensure_ascii=False),
        encoding="utf-8",
    )

    csv_path = tmp_path / "metrics.csv"
    with csv_path.open("w", encoding="utf-8", newline="") as file:
        writer = csv.writer(file)
        writer.writerow(["metric", "value"])
        writer.writerow(["teachers_total", "48"])

    docx_path = tmp_path / "acts.docx"
    document = DocxDocument()
    document.add_paragraph("На сайте опубликованы локальные нормативные акты.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Кадры"
    table.rows[0].cells[1].text = "48 преподавателей"
    document.save(docx_path)

    xlsx_path = tmp_path / "programs.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Programs"
    sheet.append(["program", "status"])
    sheet.append(["Информационные системы", "active"])
    workbook.save(xlsx_path)

    pdf_path = tmp_path / "scan.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=300, height=300)
    with pdf_path.open("wb") as file:
        writer.write(file)

    txt_result = extract_document(str(txt_path))
    json_result = extract_document(str(json_path))
    csv_result = extract_document(str(csv_path))
    docx_result = extract_document(str(docx_path))
    xlsx_result = extract_document(str(xlsx_path))
    pdf_result = extract_document(str(pdf_path))

    assert "Кадровый состав" in txt_result.text
    assert len(txt_result.fragments) == 2

    assert "accreditation" in json_result.text
    assert len(json_result.fragments) >= 1

    assert "teachers_total | 48" in csv_result.text
    assert len(csv_result.fragments) == 2

    assert "локальные нормативные акты" in docx_result.text
    assert any("48 преподавателей" in fragment.text for fragment in docx_result.fragments)

    assert "Информационные системы" in xlsx_result.text
    assert len(xlsx_result.fragments) >= 2

    assert pdf_result.requires_review is True
    assert pdf_result.page_count == 1
