from __future__ import annotations

from pathlib import Path

from docx import Document

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency in some envs
    PdfReader = None

try:
    from openpyxl import load_workbook
except Exception:  # pragma: no cover - optional dependency in some envs
    load_workbook = None


TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log"}
EXCEL_EXTENSIONS = {".xlsx", ".xlsm"}


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="cp1251", errors="ignore")


def _read_docx(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)
    return "\n".join(chunks)


def _read_pdf(path: Path) -> str:
    if PdfReader is None:
        return ""
    chunks: list[str] = []
    try:
        reader = PdfReader(str(path))
    except Exception:
        return ""
    for page in reader.pages:
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            chunks.append(text)
    return "\n".join(chunks)


def _read_excel(path: Path) -> str:
    if load_workbook is None:
        return ""
    try:
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    except Exception:
        return ""

    chunks: list[str] = []
    max_rows = 5000
    seen_rows = 0
    for sheet in workbook.worksheets:
        chunks.append(f"[Лист: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            if seen_rows >= max_rows:
                chunks.append("[Достигнут лимит строк при извлечении]")
                return "\n".join(chunks)
            values = [str(cell).strip() for cell in row if cell not in (None, "")]
            if not values:
                continue
            chunks.append(" | ".join(values))
            seen_rows += 1
    return "\n".join(chunks)


def read_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return _read_text_file(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix in EXCEL_EXTENSIONS:
        return _read_excel(path)
    return ""


def normalize_text(text: str) -> str:
    collapsed = " ".join(text.replace("\xa0", " ").split())
    return collapsed.strip()


def make_excerpt(text: str, start: int, end: int, radius: int = 100) -> str:
    left = max(0, start - radius)
    right = min(len(text), end + radius)
    excerpt = text[left:right]
    excerpt = excerpt.replace("\n", " ").replace("\r", " ")
    return normalize_text(excerpt)
